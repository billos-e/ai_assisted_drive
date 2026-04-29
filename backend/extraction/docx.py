from __future__ import annotations

from io import BytesIO

from docx import Document


def extract_docx_bytes(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_texts: list[str] = []
    for table in document.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_values:
                table_texts.append(" | ".join(row_values))
    return "\n".join(paragraphs + table_texts)
